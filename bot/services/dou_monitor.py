"""Monitor de Medidas Provisórias no Diário Oficial da União (Inlabs/DOU).

Autônomo — não compartilha código nem estado com o Monitor-de-MP externo.
Fluxo: autentica no Inlabs → baixa ZIPs do DOU (DO1E + DO1) → extrai MPs
publicadas → gera nota técnica via Claude → entrega no Telegram (mensagem
+ DOCX). Dedup por (usuário, número, ano) na tabela dou_seen_mps.

Credenciais: INLABS_EMAIL / INLABS_PASSWORD (cadastro gratuito em
inlabs.in.gov.br). Reusa ANTHROPIC_API_KEY do bot pra nota técnica.
"""
from __future__ import annotations

import asyncio
import io
import logging
import re
import zipfile
from datetime import date, timedelta
from xml.etree import ElementTree as ET

import httpx
from bs4 import BeautifulSoup
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from bot.config import settings
from bot.db.models import DouSeenMP

logger = logging.getLogger(__name__)

PLANALTO_BASE = "https://www.planalto.gov.br"
INLABS_BASE = "https://inlabs.in.gov.br"
DOU_SECTIONS = ["DO1E", "DO1"]

_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
}

_TITLE_RE = re.compile(
    r"^\s*MEDIDA PROVIS[ÓO]RIA\s+N[ºo°\.°]?\s*([\d\.]+)", re.IGNORECASE
)


class DouError(Exception):
    pass


# ──────────────────────── fetch (sync; rodar em thread) ────────────────────────

def _planalto_period(year: int) -> str:
    start = ((year - 1991) // 4) * 4 + 1991
    return f"{start}-{start + 3}"


def _fetch_mp_page(client: httpx.Client, url: str) -> tuple[str, str]:
    """Best-effort: busca ementa + texto limpo na página do Planalto."""
    try:
        resp = client.get(url, timeout=20.0)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.decompose()
        lines = [ln.strip() for ln in soup.get_text("\n", strip=True).splitlines() if ln.strip()]
        ementa = ""
        found_title = False
        for ln in lines[:40]:
            if re.match(r"MEDIDA PROVIS[ÓO]RIA\s+N", ln, re.I):
                found_title = True
                continue
            if found_title and len(ln) > 20:
                if not re.match(r"(A PRESIDENTA|O PRESIDENTE|O VICE)", ln, re.I):
                    ementa = ln
                    break
        if not ementa:
            ementa = lines[0] if lines else ""
        return ementa, "\n".join(lines[:500])
    except Exception as exc:
        logger.warning("dou: erro ao buscar página planalto %s: %s", url, exc)
        return "", ""


def _build_mp_dict(client: httpx.Client, numero: str, year: int, text_excerpt: str, target_date: date) -> dict:
    period = _planalto_period(year)
    ano2d = str(year)[-2:]
    planalto_url = f"{PLANALTO_BASE}/ccivil_03/_Ato{period}/{year}/Mpv/mpv{numero}-{ano2d}.htm"
    ementa_page, texto_planalto = _fetch_mp_page(client, planalto_url)
    if ementa_page:
        ementa = ementa_page
    else:
        stripped = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", text_excerpt)).strip()
        m = re.search(
            r"DE\s+\d{4}\s+(.+?)(?=\s+O\s+PRESIDENTE|\s+A\s+PRESIDENTA|\s+O\s+VICE)",
            stripped, re.IGNORECASE | re.DOTALL,
        )
        ementa = (re.sub(r"\s+", " ", m.group(1)).strip()[:300] if m else stripped[:300])
    return {
        "numero": numero,
        "ano": year,
        "ementa": ementa,
        "data_publicacao": target_date.isoformat(),
        "url_planalto": planalto_url,
        "texto_integral": texto_planalto or text_excerpt[:6000],
    }


def _parse_dou_xml(client: httpx.Client, xml_content: str, target_date: date) -> list[dict]:
    year = target_date.year
    results: list[dict] = []
    seen: set[str] = set()

    def _try(title_text: str, body_text: str) -> None:
        m = _TITLE_RE.match(title_text.strip().upper())
        if not m:
            return
        numero = m.group(1).replace(".", "")
        if numero in seen:
            return
        seen.add(numero)
        results.append(_build_mp_dict(client, numero, year, body_text or title_text, target_date))

    try:
        root = ET.fromstring(xml_content)
    except ET.ParseError:
        soup = BeautifulSoup(xml_content, "html.parser")
        for tag in soup.find_all(True):
            txt = tag.get_text(" ", strip=True)
            if txt:
                _try(txt, txt)
        return results

    parent_map = {child: parent for parent in root.iter() for child in parent}
    for elem in root.iter():
        text = (elem.text or "").strip()
        if not text or len(text) > 2000:
            continue
        parent = parent_map.get(elem, elem)
        body_text = ET.tostring(parent, encoding="unicode", method="text")
        _try(text, body_text)
    for elem in root.iter():
        attr_title = elem.get("title", "").strip()
        if attr_title:
            _try(attr_title, ET.tostring(elem, encoding="unicode", method="text"))

    if not results:
        for raw_line in xml_content.splitlines():
            line = re.sub(r"<[^>]+>", " ", raw_line).strip()
            if 20 <= len(line) <= 400:
                _try(line, xml_content)
    return results


def _fetch_mps_sync(target_date: date) -> list[dict]:
    email = settings.inlabs_email
    password = settings.inlabs_password.get_secret_value() if settings.inlabs_password else None
    if not email or not password:
        raise DouError(
            "INLABS_EMAIL / INLABS_PASSWORD não configurados "
            "(cadastro gratuito em inlabs.in.gov.br)."
        )

    results: list[dict] = []
    seen_numeros: set[str] = set()
    with httpx.Client(headers=_HEADERS, follow_redirects=True) as client:
        # login
        try:
            resp = client.post(
                f"{INLABS_BASE}/logar.php",
                data={"email": email, "password": password},
                headers={"Content-Type": "application/x-www-form-urlencoded"},
                timeout=20.0,
            )
            resp.raise_for_status()
        except Exception as exc:
            raise DouError(f"falha ao autenticar no Inlabs: {exc}")
        cookie = client.cookies.get("inlabs_session_cookie")
        if not cookie:
            raise DouError("autenticação Inlabs falhou (cookie ausente) — verifique e-mail/senha.")

        date_str = target_date.strftime("%Y-%m-%d")
        for section in DOU_SECTIONS:
            url = f"{INLABS_BASE}/index.php?p={date_str}&dl={date_str}-{section}.zip"
            try:
                r = client.get(
                    url, headers={"Cookie": f"inlabs_session_cookie={cookie}"}, timeout=60.0,
                )
                if r.status_code == 404:
                    continue
                r.raise_for_status()
                content = r.content
                if len(content) < 100 or content[:2] != b"PK":
                    continue  # seção não publicada / resposta HTML
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    for name in (n for n in zf.namelist() if n.lower().endswith(".xml")):
                        xml_data = zf.read(name).decode("utf-8", errors="replace")
                        if "MEDIDA PROVIS" not in xml_data.upper():
                            continue
                        for mp in _parse_dou_xml(client, xml_data, target_date):
                            if mp["numero"] not in seen_numeros:
                                seen_numeros.add(mp["numero"])
                                results.append(mp)
            except zipfile.BadZipFile:
                logger.warning("dou: ZIP inválido em %s", section)
            except Exception as exc:
                logger.warning("dou: erro processando %s: %s", section, exc)
    return results


async def fetch_mps(target_date: date) -> list[dict]:
    """Busca MPs publicadas na data (Inlabs DO1E + DO1). Roda o I/O
    bloqueante numa thread pra não travar o event loop."""
    return await asyncio.to_thread(_fetch_mps_sync, target_date)


# ──────────────────────── nota técnica (Claude) ────────────────────────

_NOTA_SYSTEM = (
    "Você é um analista legislativo sênior especializado em Medidas Provisórias "
    "do governo federal brasileiro. Redige Notas Técnicas com rigor jurídico, "
    "linguagem técnico-legislativa densa e objetiva.\n\n"
    "REGRAS:\n"
    "- Tom técnico-legislativo, objetivo — sem opiniões pessoais.\n"
    "- Leis pelo número e data: 'Lei nº 11.977, de 7 de julho de 2009'.\n"
    "- Dispositivos pelo número completo: 'art. 5º, § 1º-A, inciso II'.\n"
    "- Use travessões (—) para explicações incidentais.\n"
    "- Cite MPs correlatas pelo número e ano quando aplicável.\n"
    "- NÃO mencione quem assinou ou referendou a MP.\n\n"
    "RESUMO: um parágrafo apresentando a MP — o que faz, qual lei altera ou "
    "regime cria, e o objetivo central.\n"
    "ALTERAÇÕES: do segundo parágrafo em diante, descreva cada alteração legal "
    "(artigo/parágrafo/inciso afetado, com lei e data; efeito prático; quem é "
    "afetado; normas infralegais necessárias). Separe parágrafos com linha em "
    "branco. Não resuma em excesso."
)

_NOTA_TOOL = {
    "name": "nota_tecnica",
    "description": "Gera o conteúdo textual da Nota Técnica da Medida Provisória.",
    "input_schema": {
        "type": "object",
        "properties": {
            "titulo": {"type": "string", "description": "Nome oficial: 'Medida Provisória nº X.XXX, de AAAA'"},
            "resumo": {"type": "string", "description": "Um parágrafo resumindo a MP."},
            "alteracoes": {"type": "string", "description": "Parágrafos (\\n\\n) com as alterações legais."},
        },
        "required": ["titulo", "resumo", "alteracoes"],
    },
}


async def generate_nota_tecnica(mp: dict) -> dict | None:
    """Gera nota técnica via Claude. Retorna {titulo, resumo, alteracoes}
    ou None se a chave Anthropic não estiver configurada / falhar."""
    if not settings.anthropic_api_key:
        logger.warning("dou: ANTHROPIC_API_KEY ausente; nota técnica pulada")
        return None
    try:
        import anthropic
    except ImportError:
        return None

    user_content = (
        f"MP nº {mp['numero']}/{mp['ano']}\n"
        f"Ementa: {mp['ementa']}\n"
        f"URL: {mp.get('url_planalto', 'N/A')}\n\n"
        f"Texto integral (use para embasar a análise):\n"
        f"{(mp.get('texto_integral') or 'Não disponível')[:6000]}"
    )
    client = anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)
    try:
        resp = await client.messages.create(
            model=settings.anthropic_model,
            max_tokens=4096,
            system=[{"type": "text", "text": _NOTA_SYSTEM, "cache_control": {"type": "ephemeral"}}],
            tools=[_NOTA_TOOL],
            tool_choice={"type": "tool", "name": "nota_tecnica"},
            messages=[{"role": "user", "content": user_content}],
        )
        for block in resp.content:
            if getattr(block, "type", None) == "tool_use":
                return block.input
        return None
    except Exception:
        logger.exception("dou: falha ao gerar nota técnica MP %s/%s", mp["numero"], mp["ano"])
        return None


# ──────────────────────── prazos + formatação ────────────────────────

def compute_prazos(pub: date) -> dict:
    """Prazos regimentais a partir da publicação (dia 1 = publicação)."""
    return {
        "eficacia_fim": pub + timedelta(days=59),       # 60 dias
        "sobrestamento": pub + timedelta(days=45),
        "emendas_fim": pub + timedelta(days=7),
    }


def _br(d: date) -> str:
    return d.strftime("%d/%m/%Y")


def format_telegram_message(mp: dict, nota: dict | None) -> str:
    pub = date.fromisoformat(mp["data_publicacao"])
    prazos = compute_prazos(pub)
    titulo = (nota or {}).get("titulo") or f"Medida Provisória nº {mp['numero']}, de {mp['ano']}"
    resumo = (nota or {}).get("resumo") or mp.get("ementa") or "(sem ementa)"
    lines = [
        f"📄 <b>{titulo}</b>",
        f"<i>Publicada no DOU em {_br(pub)}</i>",
        "",
        resumo,
        "",
        "⏱️ <b>Prazos</b>",
        f"• Emendas até <b>{_br(prazos['emendas_fim'])}</b>",
        f"• Sobrestamento de pauta a partir de {_br(prazos['sobrestamento'])}",
        f"• Vigência até {_br(prazos['eficacia_fim'])} (prorrogável +60 dias)",
        "",
        f'<a href="{mp.get("url_planalto", "")}">texto no Planalto</a>',
    ]
    if nota:
        lines.append("\n📎 Nota técnica completa no anexo.")
    return "\n".join(lines)


def build_docx(mp: dict, nota: dict | None) -> bytes:
    """Monta a Nota Técnica em DOCX (sem logo/template — limpo). Retorna bytes."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    pub = date.fromisoformat(mp["data_publicacao"])
    prazos = compute_prazos(pub)
    titulo = (nota or {}).get("titulo") or f"Medida Provisória nº {mp['numero']}, de {mp['ano']}"
    resumo = (nota or {}).get("resumo") or mp.get("ementa") or "(sem ementa)"
    alteracoes = (nota or {}).get("alteracoes") or ""

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("NOTA TÉCNICA")
    r.bold = True
    r.font.size = Pt(14)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run(titulo)
    rt.bold = True
    rt.font.size = Pt(12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Publicada no DOU em {_br(pub)}").italic = True

    # Prazos
    pz = doc.add_paragraph()
    pz.add_run("Prazos regimentais").bold = True
    for label, value, red in [
        ("Emendas até", _br(prazos["emendas_fim"]), True),
        ("Sobrestamento de pauta a partir de", _br(prazos["sobrestamento"]), False),
        ("Vigência até", f"{_br(prazos['eficacia_fim'])} (prorrogável por mais 60 dias)", False),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ")
        rv = p.add_run(value)
        rv.bold = red
        if red:
            rv.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()
    rp = doc.add_paragraph().add_run("RESUMO")
    rp.bold = True
    par = doc.add_paragraph(resumo)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if alteracoes.strip():
        doc.add_paragraph()
        ra = doc.add_paragraph().add_run("ALTERAÇÕES LEGAIS")
        ra.bold = True
        for bloco in alteracoes.split("\n\n"):
            bloco = bloco.strip()
            if bloco:
                pp = doc.add_paragraph(bloco)
                pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if mp.get("url_planalto"):
        doc.add_paragraph()
        doc.add_paragraph(f"Texto integral: {mp['url_planalto']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_filename(mp: dict) -> str:
    return f"MP_{mp['numero']}_{mp['ano']}.docx"


# ──────────────────────── dedup ────────────────────────

async def filter_unseen(session: AsyncSession, user_id: int, mps: list[dict]) -> list[dict]:
    if not mps:
        return []
    rows = await session.scalars(
        select(DouSeenMP).where(DouSeenMP.user_id == user_id)
    )
    seen = {(r.numero, r.ano) for r in rows}
    return [mp for mp in mps if (mp["numero"], mp["ano"]) not in seen]


async def mark_seen(session: AsyncSession, user_id: int, mp: dict) -> None:
    session.add(DouSeenMP(user_id=user_id, numero=mp["numero"], ano=mp["ano"]))
    await session.commit()


# ──────────────────────── entrega (mensagem + DOCX) ────────────────────────

async def deliver_to_user(bot, session: AsyncSession, user, target_date: date) -> int:
    """Busca MPs novas (não vistas) da data, gera nota + DOCX e entrega no
    Telegram. Marca como vistas. Retorna quantas MPs foram entregues.
    Levanta DouError se o fetch falhar (credencial, rede)."""
    from aiogram.types import BufferedInputFile

    mps = await fetch_mps(target_date)
    novas = await filter_unseen(session, user.id, mps)
    if not novas:
        return 0

    enviadas = 0
    for mp in novas:
        nota = await generate_nota_tecnica(mp)
        try:
            await bot.send_message(
                user.id, format_telegram_message(mp, nota),
                parse_mode="HTML", disable_web_page_preview=True,
            )
            docx_bytes = await asyncio.to_thread(build_docx, mp, nota)
            await bot.send_document(
                user.id,
                BufferedInputFile(docx_bytes, filename=docx_filename(mp)),
            )
        except Exception:
            logger.exception("dou: falha ao entregar MP %s/%s ao user %s",
                             mp["numero"], mp["ano"], user.id)
            continue
        await mark_seen(session, user.id, mp)
        enviadas += 1
    return enviadas
